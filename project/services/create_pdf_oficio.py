import json
import locale
import os
from pathlib import Path
from docx import Document
from docx.shared import Pt, Inches, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH

try:
    locale.setlocale(locale.LC_ALL, 'pt_BR.UTF-8')
except locale.Error:
    """
    Se falhar, busca o nome do locale sem o '.UTF-8'
    """
    locale.setlocale(locale.LC_ALL, 'pt_BR')


def create_word_doc(file_path_docx: str, data: dict):
    """
    Gera o ofício em formato Word (.docx)
    """
    processo_num = data['numero_processo']
    entidade = data['entidade'].upper()
    cnpj = data['cnpj']
    endereco = data['endereco']
    cep = data['cep']
    uf = data['uf_proposta']
    municipio = data['municipio_proposto']
    fistel = data['fistel']
    canal= f"{data['canal_proposto']}/{data['classe_proposta']}"
    servico = data['servico']
    ipca = data['ipca']
    ipca_normalized = locale.format_string("%.2f", float(ipca), grouping=True)

    # --- Criação do Documento Word ---
    doc = Document()
    
    # Configuração da fonte padrão para o documento
    style = doc.styles['Normal']
    font = style.font
    font.name = 'Arial'
    font.size = Pt(10)
    

    # --- Bloco 1: Destinatário ---
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    p.add_run('Ao (À) Senhor(a)\n')
    p.add_run('Representante Legal da\n')
    p.add_run(f'{entidade} (CNPJ: {cnpj})\n').bold = True
    p.add_run(f'{endereco}\n')
    p.add_run(f'{cep} - {municipio}/{uf}')
    p.paragraph_format.space_after = Pt(20)

    # --- Bloco 2: Assunto ---
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    p.add_run('Assunto: ')
    p.add_run('Alteração de Plano Básico com mudança de grupo de enquadramento.').bold = True
    p.paragraph_format.space_after = Pt(6)

    # --- Bloco 3: Referência ---
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    p.add_run('Referência: ')
    p.add_run(f'Processo nº {processo_num}').bold = True
    p.paragraph_format.space_after = Pt(20)

    # --- Bloco 4: Vocativo ---
    p = doc.add_paragraph('Senhor Representante Legal,')
    p.paragraph_format.first_line_indent = Inches(1) # Recuo de 1ª linha
    p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    p.paragraph_format.space_after = Pt(12)

    # --- Bloco 5 ---
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    
    p = doc.add_paragraph()
    p.paragraph_format.tab_stops.add_tab_stop(Inches(1))
    p.add_run('1.\t')

    p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    
    p.add_run('Informamos que a Guia de Recolhimento da União ')
    p.add_run('(GRU)').bold = True
    p.add_run(' referente à cobrança sob o código de receita 6530 está disponível no Sistema Integrado de Gestão de Créditos da Anatel ')
    p.add_run('(SIGEC)').bold = True
    p.add_run('. Essa guia corresponde ao valor a ser pago pela sua entidade em função da proposta de ')
    p.add_run('promoção de classe').bold = True
    p.add_run(' do seu canal, conforme estabelecido em ')
    p.add_run('Consulta Pública.').bold = True
    p.paragraph_format.space_after = Pt(6)

    # --- Bloco 6: Texto Valor (Com recuo e negrito) ---
    p = doc.add_paragraph()
    p.paragraph_format.tab_stops.add_tab_stop(Inches(1))
    p.add_run('2.\t')

    p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

    p.add_run('O valor foi calculado com base na diferença entre os preços mínimos definidos para cada grupo de enquadramento, conforme estipulado pela ')
    p.add_run('Portaria de Consolidação GM/MCom nº 1, de 1º de junho de 2023').bold = True
    p.add_run('. A quantia indicada no quadro abaixo foi atualizada pelo ')
    p.add_run('IPCA').bold = True
    p.add_run(' desde agosto de 2013, conforme orientação do Ministério das Comunicações.')
    p.paragraph_format.space_after = Pt(18)
    p.paragraph_format.space_after = Pt(18)

    # --- TABELA ---
    table = doc.add_table(rows=2, cols=5)
    table.style = 'Table Grid'

    table.autofit = False 
    table.allow_autofit = False

    widths = (Cm(4.5), Cm(4.5), Cm(3.0), Cm(2.5), Cm(3.0))
    
    # Aplicar as larguras em cada coluna e celula
    for col_idx, width in enumerate(widths):
        table.columns[col_idx].width = width
        for cell in table.columns[col_idx].cells:
            cell.width = width

    # Cabeçalho
    hdr_cells = table.rows[0].cells
    headers = ['UF/MUNICÍPIO', 'CNPJ', 'FISTEL', 'CANAL', 'VALOR (R$)']
    for i, text in enumerate(headers):
        p = hdr_cells[i].paragraphs[0]
        run = p.add_run(text)
        run.bold = True
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER

    # Dados
    row_cells = table.rows[1].cells
    
    cnpj_formatado = cnpj.replace('-', '\u2011') 
    uf_municipio = f"{uf}/{municipio}"
    canal_classe = f"{canal} ({servico})"
    
    values = [uf_municipio, cnpj_formatado, fistel, canal_classe, ipca_normalized]
    
    for i, text in enumerate(values):
        p = row_cells[i].paragraphs[0]
        p.add_run(text)
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER

    doc.add_paragraph().paragraph_format.space_after = Pt(12)

    # --- Instruções ---
    p = doc.add_paragraph()
    p.paragraph_format.tab_stops.add_tab_stop(Inches(1))
    
    p.add_run('3.\t')
    p.add_run('Para acessar a guia de pagamento, siga as instruções abaixo:').bold = True
    
    instrucoes = [
        "a) Acesse o site http://sistemas.anatel.gov.br/boleto;",
        "b) Clique em “IMPRESSÃO DE BOLETOS”;",
        "c) Na próxima tela, insira o CNPJ da empresa e o Fistel do serviço;",
        "d) Clique em “Confirmar”;",
        "e) Na tela seguinte, mantenha selecionado “Tipo de Boleto: DEVEDORES” e “Consultar por: CNPJ/CPF”;",
        "f) Clique novamente em “Confirmar”;",
        "g) Para imprimir a GRU, clique no ícone da impressora."
    ]

    for item in instrucoes:
        p = doc.add_paragraph(item)
        p.paragraph_format.left_indent = Inches(1.2)
        p.paragraph_format.space_after = Pt(2)

    # --- Bloco Final ---
    p = doc.add_paragraph()
    p.paragraph_format.tab_stops.add_tab_stop(Inches(1))
    p.add_run('4.\t')
    p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    p.paragraph_format.space_before = Pt(12) 
    p.add_run('Conforme o ')
    p.add_run('art. 31, §1º, da Portaria GM/MCom nº 1/2023').bold = True
    p.add_run(', nos casos em que é exigido o pagamento pela diferença entre os preços mínimos de outorga, a Anatel somente procederá com a alteração no respectivo Plano Básico e autorizará as novas condições de operação após a confirmação do pagamento.')
    
    # Atenciosamente
    p = doc.add_paragraph('Atenciosamente,')
    p.paragraph_format.first_line_indent = Inches(1)
    p.paragraph_format.space_before = Pt(20)

    # --- Salvar ---
    doc.save(file_path_docx)